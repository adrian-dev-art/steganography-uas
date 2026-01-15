import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_formatted_text(paragraph, text):
    # Basic bold/italic/inline-code parser
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(part)

        # Apply Times New Roman and Black color unless it's code
        if not run.font.name == 'Courier New':
            run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)

def md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()

    # Set default style to Times New Roman and Black
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False

    for line in lines:
        line = line.rstrip()

        # Code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            run = p.runs[0] if p.runs else p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 0, 0)
            continue

        # Images - Robust detection
        image_match = re.search(r'!\[(.*?)\]\((.*?)\)', line)
        if image_match:
            img_path = image_match.group(2)
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Pt(400))
                    last_p = doc.paragraphs[-1]
                    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    p = doc.add_paragraph(f"[Error loading image: {img_path}]")
                    p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            else:
                p = doc.add_paragraph(f"[Image not found: {img_path}]")
                p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            continue

        # Headers
        header_match = re.match(r'^(#+)\s+(.*)', line)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2)
            h = doc.add_heading(text, level=min(level, 9))
            for run in h.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)
            continue

        # Horizontal rules
        if re.match(r'^---+$', line):
            doc.add_page_break()
            continue

        # Lists
        list_match = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if list_match:
            text = list_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            continue

        num_list_match = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if num_list_match:
            text = num_list_match.group(2)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            continue

        # Blockquotes (Alerts)
        if line.startswith('>'):
            text = line.lstrip('> ').strip()
            if text.startswith('[!'):
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(20)
            run = p.add_run(text)
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
            continue

        # Tables
        if line.startswith('|'):
            if '---' in line:
                continue # Skip alignment row

            cells = [c.strip() for c in line.split('|') if c.strip()]
            if not cells:
                continue

            # If it's the first row of a table, create a new table
            if not hasattr(doc, '_current_table') or doc._current_table is None:
                doc._current_table = doc.add_table(rows=0, cols=len(cells))
                doc._current_table.style = 'Table Grid'

            row_cells = doc._current_table.add_row().cells
            for i, cell_text in enumerate(cells):
                p = row_cells[i].paragraphs[0]
                add_formatted_text(p, cell_text)
            continue
        else:
            doc._current_table = None # Reset table tracking

        # Regular paragraphs
        if line.strip():
            p = doc.add_paragraph()
            add_formatted_text(p, line)

    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

if __name__ == "__main__":
    output_dir = "format laporan"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    md_to_docx("Langkah_Investigasi_Extracting.md", "Langkah_Investigasi_Extracting.docx")
    md_to_docx("Laporan_Hasil_Investigasi_Digital_Forensik.md", "Laporan_Hasil_Investigasi_Digital_Forensik.docx")
    md_to_docx("Laporan_Forensik_Suspect_Image.md", "Laporan_Forensik_Suspect_Image.docx")
    md_to_docx("README.md", os.path.join(output_dir, "README.docx"))
